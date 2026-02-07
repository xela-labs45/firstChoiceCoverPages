from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_border(doc):
    """Adds a box border to the first section of the document."""
    sec_pr = doc.sections[0]._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')  # Border size (1/8 pt)
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '003366') # Dark Blue
        pg_borders.append(border)
    
    sec_pr.append(pg_borders)

def create_template():
    doc = Document()
    
    # 1. Add Page Border
    add_page_border(doc)

    # 2. Add School Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ACADEMIC ASSESSMENT COVER")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

    # Spacer
    doc.add_paragraph()

    # 3. Add Subject Placeholder (The Big Title)
    subject_para = doc.add_paragraph()
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subject_para.add_run("{{Subject}}")
    run_sub.bold = True
    run_sub.font.size = Pt(36)
    run_sub.font.color.rgb = RGBColor(0, 0, 0)
    
    # Spacer
    doc.add_paragraph()
    doc.add_paragraph()

    # 4. Student Details Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid' # Standard grid look
    table.autofit = False
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.0)

    # Row 1: Name
    cell = table.cell(0, 0)
    cell.text = "Student Name:"
    cell.paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = "{{Name}}"

    # Row 2: Surname
    cell = table.cell(1, 0)
    cell.text = "Student Surname:"
    cell.paragraphs[0].runs[0].bold = True
    table.cell(1, 1).text = "{{Surname}}"

    # Row 3: Class
    cell = table.cell(2, 0)
    cell.text = "Class:"
    cell.paragraphs[0].runs[0].bold = True
    table.cell(2, 1).text = "{{Class}}"

    # Row 4: Year
    cell = table.cell(3, 0)
    cell.text = "Academic Year:"
    cell.paragraphs[0].runs[0].bold = True
    table.cell(3, 1).text = "{{Year}}"

    # Spacer
    doc.add_paragraph()
    doc.add_paragraph()

    # 5. Teacher/Grades Section (Static area for writing)
    comment_header = doc.add_paragraph()
    run_comment = comment_header.add_run("Teacher's Comments / Grade:")
    run_comment.bold = True
    run_comment.font.size = Pt(14)
    run_comment.font.underline = True

    # Create a box for comments (using a single cell table)
    comment_table = doc.add_table(rows=1, cols=1)
    comment_table.style = 'Table Grid'
    comment_cell = comment_table.cell(0, 0)
    comment_cell.height = Inches(2.5) # Make it tall
    
    # Save the file
    file_name = "template.docx"
    doc.save(file_name)
    print(f"Success! '{file_name}' has been created.")

if __name__ == "__main__":
    create_template()