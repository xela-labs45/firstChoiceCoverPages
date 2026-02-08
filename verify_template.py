from docx import Document
import os

def check_placeholders(file_path):
    if not os.path.exists(file_path):
        print(f"File {file_path} does not exist.")
        return

    doc = Document(file_path)
    print(f"Checking {file_path} for placeholders...")
    
    placeholders = ["{{Name}}", "{{Surname}}", "{{Class}}", "{{Year}}", "{{Subject}}"]
    found = {p: False for p in placeholders}
    
    # Check paragraphs
    for p in doc.paragraphs:
        for placeholder in placeholders:
            if placeholder in p.text:
                found[placeholder] = True
                
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for placeholder in placeholders:
                        if placeholder in p.text:
                            found[placeholder] = True
                            
    for p, is_found in found.items():
        status = "✅ FOUND" if is_found else "❌ MISSING"
        print(f"{p}: {status}")

if __name__ == "__main__":
    check_placeholders("template.docx")
