import streamlit as st
import os
import io
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt
from datetime import datetime

# First Choice Student Cover Page Generator
# =========================================
#
# This application generates professional student cover pages from a Word (.docx) template.
# It allows for batch generation where multiple subjects are merged into a single 
# printable Word document, preserving formatting and layout for each page.
#
# TEMPLATE REQUIREMENTS:
# ----------------------
# The application requires a 'template.docx' file in the same directory.
# The template must contain the following placeholders exactly as shown:
# - {{Name}}    : The student's first name
# - {{Surname}} : The student's surname
# - {{Class}}   : The student's class (e.g., Grade 10)
# - {{Year}}    : The academic year
# - {{Subject}} : The specific subject (one page is generated per subject)
#
# Note: The script is designed to preserve the font, size, and styling of these
# placeholders as defined in your Word template.

# --- Configuration & Setup ---
st.set_page_config(page_title="Student Cover Page Generator", layout="wide")

# --- Helper Functions ---

def replace_placeholder(doc, placeholder, replacement):
    """
    Replaces a placeholder in a python-docx Document object.
    
    ROBUST REPLACEMENT LOGIC:
    1. Tries to replace text at the 'Run' level to keep specific character formatting.
    2. If the placeholder is 'split' across multiple runs (common in Word), it 
       falls back to Paragraph-level replacement.
    3. During fallback, it captures the font name, size, bold, italic, and color 
       from the first run and re-applies it to the new text to ensure the 
       cover page still looks exactly like the template.
    """
    replacement = str(replacement)
    
    def replace_in_element(element):
        for paragraph in element.paragraphs:
            if placeholder in paragraph.text:
                # Attempt 1: Run-level replacement (preserves specific formatting)
                replaced_in_run = False
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
                        replaced_in_run = True
                
                # Attempt 2: Paragraph-level fallback (handles placeholders split by Word)
                if not replaced_in_run:
                    # Capture formatting from the existing text before it's wiped
                    style_props = {
                        "name": None, "size": None, "bold": None, 
                        "italic": None, "underline": None, "color": None
                    }
                    
                    if paragraph.runs:
                        r = paragraph.runs[0] # Carry the main style from the first run
                        style_props["name"] = r.font.name
                        style_props["size"] = r.font.size
                        style_props["bold"] = r.bold
                        style_props["italic"] = r.italic
                        style_props["underline"] = r.underline
                        try:
                            style_props["color"] = r.font.color.rgb
                        except:
                            style_props["color"] = None
                    
                    # Replace the entire paragraph text
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                    
                    # Re-apply the captured formatting to the new run
                    if paragraph.runs:
                        new_run = paragraph.runs[0]
                        if style_props["name"]: new_run.font.name = style_props["name"]
                        if style_props["size"]: new_run.font.size = style_props["size"]
                        if style_props["bold"] is not None: new_run.bold = style_props["bold"]
                        if style_props["italic"] is not None: new_run.italic = style_props["italic"]
                        if style_props["underline"] is not None: new_run.underline = style_props["underline"]
                        if style_props["color"]: new_run.font.color.rgb = style_props["color"]

    # Replace in main body
    replace_in_element(doc)

    # Replace in all tables (where headers/details often live)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_element(cell)

def generate_single_document(template_path, student_data, subjects):
    """
    Generates a single Word document containing multiple cover pages.
    
    MERGING STRATEGY:
    - It uses 'docxcompose' to handle the heavy lifting of joining documents.
    - An explicit page break is added between documents to ensure each 
      cover page starts on its own sheet.
    """
    master_doc = None
    composer = None

    for i, subject in enumerate(subjects):
        # Create a fresh copy of the template for this specific subject
        temp_doc = Document(template_path)
        
        # Mapping data to placeholders
        replacements = {
            "{{Name}}": student_data.get("Name", ""),
            "{{Surname}}": student_data.get("Surname", ""),
            "{{Class}}": student_data.get("Class", ""),
            "{{Year}}": student_data.get("Year", ""),
            "{{Subject}}": subject
        }
        
        # Process replacements for this individual page
        for placeholder, value in replacements.items():
            replace_placeholder(temp_doc, placeholder, value)
            
        if master_doc is None:
            # The first document becomes the 'base' for the file
            master_doc = temp_doc
            composer = Composer(master_doc)
        else:
            # Force a new page and append the next subject's cover
            master_doc.add_page_break()
            composer.append(temp_doc)

    # Save the final merged document to an in-memory buffer
    doc_buffer = io.BytesIO()
    if master_doc:
        master_doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- Streamlit UI Logic ---

def main():
    st.title("🎓 First Choice Student Cover Page Generator")
    st.markdown("""
    This tool helps you quickly generate personalized cover pages for all your subjects at once.
    """)

    # --- Sidebar: User Inputs ---
    with st.sidebar:
        st.header("Student Details")
        name = st.text_input("Name", placeholder="First Name")
        surname = st.text_input("Surname", placeholder="Surname")
        student_class = st.text_input("Class", placeholder="e.g., Grade 10A")
        
        current_year = datetime.now().year
        year = st.number_input("Year", min_value=2000, max_value=2100, value=current_year, step=1)

    # --- Main Area: Preparation ---
    
    st.subheader("1. Template Status")
    
    # Verify the School Standard Template is present
    standard_template_path = "template.docx"
    has_standard_template = os.path.exists(standard_template_path)
    
    if has_standard_template:
        st.success("✅ School Standard Template ('template.docx') detected.")
        template_file = standard_template_path
    else:
        st.error("❌ ERROR: 'template.docx' not found in the app folder.")
        st.info("To use this app, please place your Word template named 'template.docx' in the same folder as this script.")
        template_file = None

    st.subheader("2. Subject Selection")
    
    # Pre-defined list of common subjects
    DEFAULT_SUBJECTS = [
        "Mathematics", "Science", "English", "History", "Geography", 
        "Art", "Physics", "Chemistry", "Biology", "Computer Science",
        "Physical Education", "Music", "Drama", "Economics"
    ]
    
    selected_subjects = st.multiselect("Select subjects for the cover pages:", options=DEFAULT_SUBJECTS)
    
    # Support for adding custom subjects not in the list
    custom_subjects_input = st.text_input("Add any other subjects (separate with commas):", placeholder="Robotics, Music Theory")
    if custom_subjects_input:
        custom_list = [s.strip() for s in custom_subjects_input.split(",") if s.strip()]
        selected_subjects.extend(custom_list)
        selected_subjects = list(set(selected_subjects)) # Remove duplicates

    st.subheader("3. Final Step")
    
    if st.button("Generate & Prepare Download", type="primary"):
        # Basic validation to prevent errors
        if not template_file:
            st.error("Cannot proceed without 'template.docx'.")
            return
            
        if not name or not surname:
            st.warning("Please enter the student's Name and Surname.")
            return
            
        if not selected_subjects:
            st.warning("Please select at least one subject.")
            return

        with st.spinner("Generating your document..."):
            try:
                student_data = {
                    "Name": name,
                    "Surname": surname,
                    "Class": student_class,
                    "Year": year
                }
                
                # Combine all pages into one file
                doc_buffer = generate_single_document(template_file, student_data, selected_subjects)
                
                st.success(f"Generated {len(selected_subjects)} cover pages!")
                
                # Filename sanitization for the download
                clean_name = "".join([c for c in f"{name}_{surname}" if c.isalnum() or c in (' ', '_', '-')]).replace(' ', '_')
                clean_class = "".join([c for c in student_class if c.isalnum() or c in (' ', '_', '-')]).replace(' ', '_')
                file_name = f"{clean_name}_{clean_class}_Covers.docx"
                
                st.download_button(
                    label="📥 Download all Cover Pages",
                    data=doc_buffer,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Something went wrong: {e}")
                # Optional: st.write(traceback.format_exc()) for debugging

if __name__ == "__main__":
    main()